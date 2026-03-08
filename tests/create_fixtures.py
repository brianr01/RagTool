"""Generate binary test fixtures (PDF, DOCX) that can't be created as plain text."""
import os
import sys

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def create_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    path = os.path.join(FIXTURES_DIR, "sample.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(72, 700, "Page 1: Introduction to Vector Databases")
    c.drawString(72, 680, "Vector databases store high-dimensional embeddings for similarity search.")
    c.showPage()
    c.drawString(72, 700, "Page 2: RAG Architecture")
    c.drawString(72, 680, "Retrieval-Augmented Generation combines retrieval with generation.")
    c.showPage()
    c.save()
    print(f"Created {path}")


def create_docx():
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "sample.docx")
    doc = Document()
    doc.add_paragraph("Sample Document for Testing")
    doc.add_paragraph(
        "This is a test document created for the RAG ingestion pipeline. "
        "It contains multiple paragraphs to verify DOCX extraction."
    )
    doc.add_paragraph(
        "Vector search enables semantic matching between queries and documents "
        "using cosine similarity on embedding vectors."
    )
    doc.save(path)
    print(f"Created {path}")


def create_large_pdf(num_pages=100):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    path = os.path.join(FIXTURES_DIR, "large.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    for i in range(1, num_pages + 1):
        c.drawString(72, 700, f"Page {i} of {num_pages}: Large PDF stress test")
        c.drawString(72, 680, f"This is page {i} of a {num_pages}-page document used to verify")
        c.drawString(72, 660, "that the ingestion pipeline handles large PDFs correctly.")
        c.drawString(72, 640, f"Unique content for page {i}: embedding vector similarity search RAG.")
        c.showPage()
    c.save()
    print(f"Created {path}")


if __name__ == "__main__":
    os.makedirs(FIXTURES_DIR, exist_ok=True)
    create_pdf()
    create_docx()
    create_large_pdf()
    print("All fixtures created.")
